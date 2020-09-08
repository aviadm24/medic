from PyPDF2 import PdfFileMerger,PdfFileReader,PdfFileWriter
import io
#from reportlab.pdfgen import canvas
import os
from docx import Document
from docx.styles.style import WD_STYLE_TYPE
#from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
try:
    from win32com import client
    import pythoncom
except:
    pass
import threading
#from queue import Queue
import time


def convert_to_pdf(doc):
    pythoncom.CoInitialize()
    word = client.DispatchEx("Word.Application")
    new_name = doc.replace(".docx", r".pdf")
    worddoc = word.Documents.Open(doc)
    worddoc.SaveAs(new_name, FileFormat=17)
    worddoc.Close()
    word.Quit()


def makeToc(file_name, cat_dict):
    document = Document('static\\table.docx')
    style1 = document.styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
    font = style1.font
    font.name = 'Gisha'
    # font.bold = True
    font.rtl = True
    font.color.rgb = RGBColor(3, 144, 207)
    for cateorical_dose in cat_dict.keys():
        p = document.add_paragraph(cateorical_dose, style='rtl')
        p.alignment = 2
        for line in cat_dict[cateorical_dose]:
            # text = i[0] + '....' + str(i[1])
            p = document.add_paragraph(line, style='rtl')
            # run = p.add_run(line)
            # run.style = style1
            # font = run.font
            # font.rtl = True

    r = p.add_run()
    # r.add_text('Good Morning every body,This is my ')
    r.add_picture('static\\logo.png')
    # r.add_text(' do you like it?')
    name = 'temp\\{}.docx'.format(file_name)
    document.save(name)
    path = os.getcwd() + '\\' + name
    print(path)
    convert_to_pdf(path)


# makeToc("test", )